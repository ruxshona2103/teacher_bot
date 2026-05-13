import os
import tempfile
from io import BytesIO

from aiogram import Router, F, Bot
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import (
    Message, CallbackQuery,
    InlineKeyboardMarkup, InlineKeyboardButton,
    ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove,
    BufferedInputFile,
)
from sqlalchemy import select, func, delete
from docx import Document

from db import Session
from models.models import User, Topic, Question, UserResult, UserProgress
from services.file_service import extract_text, parse_questions_with_regex, is_standard_format, count_expected_questions
from services.ai_service import parse_questions_with_ai
from config import ADMIN_IDS

router = Router()

# ── Klaviaturalar ─────────────────────────────────────────────────────────────

ADMIN_KB = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="📤 Test yuklash"),    KeyboardButton(text="📋 Testlar ro'yxati")],
    [KeyboardButton(text="📊 Natijalar"),       KeyboardButton(text="📥 Word eksport")],
], resize_keyboard=True)

UPLOAD_CONTROL_KB = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="⬅️ Orqaga"),          KeyboardButton(text="❌ Bekor qilish")],
    [KeyboardButton(text="📎 Boshqa fayl yuklash"), KeyboardButton(text="🗑 Tanlangan testni o'chir")],
], resize_keyboard=True)


def is_admin(user_id: int) -> bool:
    return user_id in ADMIN_IDS


# ── Admin panel (reply keyboard handler) ─────────────────────────────────────

@router.message(Command("admin"))
async def admin_panel(message: Message):
    if not is_admin(message.from_user.id):
        return
    await message.answer("👨‍💼 Admin panel:", reply_markup=ADMIN_KB)


@router.message(F.text == "📋 Testlar ro'yxati")
async def handle_tests_list(message: Message):
    if not is_admin(message.from_user.id):
        return
    await _show_tests_list(message)


@router.message(F.text == "📊 Natijalar")
async def handle_results(message: Message):
    if not is_admin(message.from_user.id):
        return
    await _show_results(message)


@router.message(F.text == "📥 Word eksport")
async def handle_export(message: Message):
    if not is_admin(message.from_user.id):
        return
    await _export_word(message)


# ── Test yuklash ──────────────────────────────────────────────────────────────

class UploadState(StatesGroup):
    choosing_type      = State()
    choosing_topic     = State()
    waiting_file       = State()
    confirming_answers = State()
    adding_images      = State()


@router.message(F.text == "📤 Test yuklash")
async def handle_upload_btn(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await _start_upload(message, state)


async def _start_upload(message: Message, state: FSMContext):
    await state.set_state(UploadState.choosing_type)
    # Control keyboard pastda doim ko'rinadi
    await message.answer("📤 Test yuklash:", reply_markup=UPLOAD_CONTROL_KB)
    keyboard = InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="📖 Nazariy", callback_data="upload_type:nazariy"),
        InlineKeyboardButton(text="🔧 Amaliy",  callback_data="upload_type:amaliy"),
    ]])
    await message.answer("Qaysi test turini yuklaysiz?", reply_markup=keyboard)


# ── Upload oqimi ichidagi control buttonlar ───────────────────────────────────

@router.message(F.text == "❌ Bekor qilish")
async def cancel_upload(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await state.clear()
    await message.answer("❌ Bekor qilindi.", reply_markup=ADMIN_KB)


@router.message(F.text == "⬅️ Orqaga")
async def go_back(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    current = await state.get_state()

    if current == UploadState.choosing_topic:
        await _start_upload(message, state)
    elif current == UploadState.waiting_file:
        data = await state.get_data()
        await state.set_state(UploadState.choosing_topic)
        await _show_topic_buttons(message, state)
    elif current == UploadState.confirming_answers:
        await state.set_state(UploadState.waiting_file)
        await message.answer("📎 Boshqa fayl yuboring:")
    else:
        await message.answer("👨‍💼 Admin panel:", reply_markup=ADMIN_KB)


@router.message(F.text == "📎 Boshqa fayl yuklash")
async def upload_another(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    data = await state.get_data()
    await state.update_data(questions=None, correct_answers={}, current_q=0)
    await state.set_state(UploadState.waiting_file)
    await message.answer("📎 Yangi fayl yuboring:")


@router.message(F.text == "🗑 Tanlangan testni o'chir")
async def delete_selected_topic(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    data = await state.get_data()
    topic_id = data.get("topic_id")
    if not topic_id:
        await message.answer("❌ Hech qanday test tanlanmagan.")
        return

    async with Session() as session:
        topic = await session.get(Topic, topic_id)
        if topic:
            topic.is_active = False
            await session.commit()
            await message.answer(f"🗑 '{topic.title}' o'chirildi.")
        else:
            await message.answer("❌ Test topilmadi.")

    await state.clear()
    await message.answer("👨‍💼 Admin panel:", reply_markup=ADMIN_KB)


# ── Upload oqimi ──────────────────────────────────────────────────────────────

@router.callback_query(F.data.startswith("upload_type:"), UploadState.choosing_type)
async def ask_topic(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        return
    test_type = callback.data.split(":")[1]
    await state.update_data(test_type=test_type)
    await state.set_state(UploadState.choosing_topic)
    await _show_topic_buttons(callback.message, state)


async def _show_topic_buttons(message: Message, state: FSMContext):
    data = await state.get_data()
    test_type = data["test_type"]
    type_label = "📖 Nazariy" if test_type == "nazariy" else "🔧 Amaliy"

    async with Session() as session:
        all_topics = (await session.scalars(
            select(Topic).where(Topic.is_active == True).order_by(Topic.order_num)
        )).all()

        if test_type == "amaliy":
            # Faqat Nazariy savollari mavjud bo'lgan topiclarni ko'rsatamiz
            filtered = []
            for t in all_topics:
                naz_count = await session.scalar(
                    select(func.count()).select_from(Question).where(
                        Question.topic_id == t.id,
                        Question.type == "nazariy",
                    )
                )
                if naz_count > 0:
                    filtered.append(t)
            topics = filtered
        else:
            topics = all_topics

    buttons = [[InlineKeyboardButton(text=f"📚 {t.title}", callback_data=f"upload_topic:{t.id}")]
               for t in topics]

    if test_type == "nazariy":
        buttons.append([InlineKeyboardButton(text="➕ Yangi dars qo'shish", callback_data="upload_topic:new")])

    if not buttons:
        await message.answer(
            "⚠️ Hali hech bir darsga Nazariy test qo'shilmagan.\n"
            "Avval Nazariy test yuklang."
        )
        await state.clear()
        await message.answer("👨‍💼 Admin panel:", reply_markup=ADMIN_KB)
        return

    await message.answer(
        f"{type_label} testi qaysi darsga qo'shilsin?",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons),
    )


@router.callback_query(F.data.startswith("upload_topic:"), UploadState.choosing_topic)
async def topic_chosen(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        return

    value = callback.data.split(":")[1]
    if value == "new":
        await state.update_data(topic_id=None, waiting_new_topic_name=True)
        await state.set_state(UploadState.waiting_file)
        await callback.message.answer("Yangi dars nomini kiriting (masalan: Dars 3):")
    else:
        await state.update_data(topic_id=int(value), waiting_new_topic_name=False)
        await state.set_state(UploadState.waiting_file)
        await callback.message.answer("📎 .docx, .pdf yoki .txt fayl yuboring:")


@router.message(UploadState.waiting_file, F.text)
async def receive_new_topic_name(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    if message.text in ("⬅️ Orqaga", "❌ Bekor qilish", "📎 Boshqa fayl yuklash", "🗑 Tanlangan testni o'chir"):
        return
    data = await state.get_data()
    if not data.get("waiting_new_topic_name"):
        return

    topic_name = message.text.strip()
    test_type = data["test_type"]

    async with Session() as session:
        existing_topic = await session.scalar(
            select(Topic).where(Topic.title == topic_name, Topic.is_active == True)
        )
        if existing_topic:
            q_count = await session.scalar(
                select(func.count()).select_from(Question).where(
                    Question.topic_id == existing_topic.id,
                    Question.type == test_type,
                )
            )
            type_label = "Nazariy" if test_type == "nazariy" else "Amaliy"
            if q_count > 0:
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="♻️ Ustiga yoz (almashtir)", callback_data=f"overwrite:{existing_topic.id}")],
                    [InlineKeyboardButton(text="✏️ Boshqa nom kiritaman", callback_data="rename_topic")],
                ])
                await message.answer(
                    f"⚠️ <b>'{topic_name}'</b> darsining <b>{type_label}</b> testi allaqachon mavjud!\n"
                    f"({q_count} ta savol bor)\n\nNima qilmoqchisiz?",
                    reply_markup=keyboard,
                    parse_mode="HTML",
                )
                await state.update_data(new_topic_name=topic_name, waiting_new_topic_name=False)
                return
            else:
                # Topic bor lekin bu tur uchun savol yo'q — mavjud topicga qo'shamiz
                await state.update_data(
                    topic_id=existing_topic.id,
                    new_topic_name=topic_name,
                    waiting_new_topic_name=False,
                )
                await message.answer(
                    f"✅ '{topic_name}' topildi. {type_label} testi qo'shiladi.\n📎 Fayl yuboring:"
                )
                return

    await state.update_data(new_topic_name=topic_name, waiting_new_topic_name=False)
    await message.answer("📎 Endi .docx, .pdf yoki .txt fayl yuboring:")


@router.callback_query(F.data.startswith("overwrite:"), UploadState.waiting_file)
async def overwrite_topic(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        return
    topic_id = int(callback.data.split(":")[1])
    data = await state.get_data()
    test_type = data["test_type"]

    async with Session() as session:
        await session.execute(
            delete(Question).where(
                Question.topic_id == topic_id,
                Question.type == test_type,
            )
        )
        await session.commit()

    await state.update_data(topic_id=topic_id, waiting_new_topic_name=False)
    await callback.message.answer("🗑 Eski savollar o'chirildi. 📎 Yangi fayl yuboring:")


@router.callback_query(F.data == "rename_topic", UploadState.waiting_file)
async def rename_topic(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        return
    await state.update_data(waiting_new_topic_name=True, topic_id=None)
    await callback.message.answer("✏️ Yangi dars nomini kiriting:")


@router.message(UploadState.waiting_file, F.document)
async def receive_file(message: Message, state: FSMContext, bot: Bot):
    if not is_admin(message.from_user.id):
        return

    doc = message.document
    ext = os.path.splitext(doc.file_name)[1].lower()
    if ext not in (".docx", ".pdf", ".txt"):
        await message.answer("❌ Faqat .docx, .pdf, .txt formatlar qabul qilinadi.")
        return

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
        await bot.download(doc, destination=tmp.name)
        tmp_path = tmp.name

    try:
        raw_text = extract_text(tmp_path)
    except Exception as e:
        await message.answer(f"❌ Fayl o'qishda xato: {e}")
        return
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

    expected = count_expected_questions(raw_text)

    if is_standard_format(raw_text):
        questions = parse_questions_with_regex(raw_text)
        # Regex noto'g'ri hisoblagan bo'lsa — AI bilan qayta parse
        if expected > 0 and len(questions) < expected:
            await message.answer(
                f"⚠️ Regex {len(questions)} ta topdi, lekin faylda {expected} ta savol bor. "
                f"AI bilan qayta parse qilinmoqda..."
            )
            questions = await parse_questions_with_ai(raw_text)
            method = f"ai (regex {len(questions)}/{expected} topgan)"
        else:
            method = "regex"
    else:
        await message.answer("⏳ Fayl tartibsiz format. AI bilan parse qilinmoqda...")
        questions = await parse_questions_with_ai(raw_text)
        method = "ai"

    if not questions:
        await message.answer("❌ Savollar topilmadi. Fayl formatini tekshiring.")
        await state.clear()
        await message.answer("👨‍💼 Admin panel:", reply_markup=ADMIN_KB)
        return

    await state.update_data(questions=questions, method=method, current_q=0, correct_answers={})
    await state.set_state(UploadState.confirming_answers)
    await message.answer(f"✅ {len(questions)} ta savol topildi ({method}).\n\nHar bir savolga to'g'ri javobni belgilang:")
    await _ask_correct_answer(message, state)


async def _ask_correct_answer(message: Message, state: FSMContext):
    data = await state.get_data()
    questions = data["questions"]
    current = data["current_q"]

    if current >= len(questions):
        await _save_test(message, state)
        return

    q = questions[current]
    keyboard = InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="A", callback_data="correct:A"),
        InlineKeyboardButton(text="B", callback_data="correct:B"),
        InlineKeyboardButton(text="C", callback_data="correct:C"),
        InlineKeyboardButton(text="D", callback_data="correct:D"),
    ]])

    await message.answer(
        f"Savol {current + 1}/{len(questions)}:\n\n"
        f"{q['text']}\n\n"
        f"A) {q['option_a']}\n"
        f"B) {q['option_b']}\n"
        f"C) {q['option_c']}\n"
        f"D) {q['option_d']}\n\n"
        f"To'g'ri javobni belgilang:",
        reply_markup=keyboard,
    )


@router.callback_query(F.data.startswith("correct:"), UploadState.confirming_answers)
async def receive_correct_answer(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        return

    answer = callback.data.split(":")[1]
    data = await state.get_data()
    correct_answers = data["correct_answers"]
    current = data["current_q"]

    correct_answers[str(current)] = answer
    await state.update_data(correct_answers=correct_answers, current_q=current + 1)
    await callback.answer(f"✅ {answer} belgilandi")
    await callback.message.delete()
    await _ask_correct_answer(callback.message, state)


async def _save_test(message: Message, state: FSMContext):
    data = await state.get_data()
    await state.clear()

    async with Session() as session:
        topic_id = data.get("topic_id")

        if topic_id:
            topic = await session.get(Topic, topic_id)
        else:
            existing = (await session.scalars(select(Topic))).all()
            max_order = max((t.order_num for t in existing), default=0)
            topic = Topic(
                title=data.get("new_topic_name", "Yangi dars"),
                order_num=max_order + 1,
            )
            session.add(topic)
            await session.flush()

        for i, q_data in enumerate(data["questions"]):
            session.add(Question(
                topic_id=topic.id,
                type=data["test_type"],
                text=q_data["text"],
                option_a=q_data["option_a"],
                option_b=q_data["option_b"],
                option_c=q_data["option_c"],
                option_d=q_data["option_d"],
                correct=data["correct_answers"].get(str(i), "A"),
            ))

        await session.commit()
        topic_title = topic.title
        saved_topic_id = topic.id

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📷 Ha, rasm qo'shaman", callback_data=f"add_images:{saved_topic_id}")],
        [InlineKeyboardButton(text="⏭ Yo'q, o'tkazib yuboraman", callback_data="skip_images")],
    ])
    await message.answer(
        f"✅ Test saqlandi!\n"
        f"📚 Dars: {topic_title}\n"
        f"📝 Tur: {data['test_type']}\n"
        f"❓ Savollar: {len(data['questions'])} ta\n\n"
        f"📷 Savollar uchun rasm qo'shmoqchimisiz?",
        reply_markup=keyboard,
    )


# ── Rasm qo'shish oqimi ───────────────────────────────────────────────────────

@router.callback_query(F.data == "skip_images")
async def skip_images(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        return
    await callback.message.answer("👨‍💼 Admin panel:", reply_markup=ADMIN_KB)


@router.callback_query(F.data.startswith("add_images:"))
async def start_adding_images(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        return

    topic_id = int(callback.data.split(":")[1])
    async with Session() as session:
        questions = (await session.scalars(
            select(Question).where(Question.topic_id == topic_id).order_by(Question.id)
        )).all()
        q_ids = [q.id for q in questions]

    await state.set_state(UploadState.adding_images)
    await state.set_data({"image_q_ids": q_ids, "image_current": 0})
    await _ask_question_image(callback.message, state)


async def _ask_question_image(message: Message, state: FSMContext):
    data = await state.get_data()
    q_ids = data["image_q_ids"]
    current = data["image_current"]

    if current >= len(q_ids):
        await state.clear()
        await message.answer("✅ Rasmlar saqlandi!", reply_markup=ADMIN_KB)
        return

    async with Session() as session:
        q = await session.get(Question, q_ids[current])

    keyboard = InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="⏭ Rasm kerak emas", callback_data="img_skip"),
        InlineKeyboardButton(text="⏹ To'xtatish", callback_data="img_stop"),
    ]])
    await message.answer(
        f"📷 Savol {current + 1}/{len(q_ids)}:\n\n{q.text}\n\nRasm yuboring yoki o'tkazib yuboring:",
        reply_markup=keyboard,
    )


@router.message(UploadState.adding_images, F.photo)
async def receive_question_image(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return

    data = await state.get_data()
    q_ids = data["image_q_ids"]
    current = data["image_current"]
    file_id = message.photo[-1].file_id  # eng yuqori sifatli rasm

    async with Session() as session:
        q = await session.get(Question, q_ids[current])
        q.image_file_id = file_id
        await session.commit()

    await state.update_data(image_current=current + 1)
    await message.answer("✅ Rasm saqlandi.")
    await _ask_question_image(message, state)


@router.callback_query(F.data == "img_skip", UploadState.adding_images)
async def img_skip(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        return
    data = await state.get_data()
    await state.update_data(image_current=data["image_current"] + 1)
    await _ask_question_image(callback.message, state)


@router.callback_query(F.data == "img_stop", UploadState.adding_images)
async def img_stop(callback: CallbackQuery, state: FSMContext):
    if not is_admin(callback.from_user.id):
        return
    await state.clear()
    await callback.message.answer("✅ Rasm qo'shish to'xtatildi.", reply_markup=ADMIN_KB)


# ── Testlar ro'yxati ──────────────────────────────────────────────────────────

async def _show_tests_list(message: Message):
    async with Session() as session:
        topics = (await session.scalars(
            select(Topic).where(Topic.is_active == True).order_by(Topic.order_num)
        )).all()

    if not topics:
        await message.answer("📭 Hech qanday test yo'q.")
        return

    buttons = [[InlineKeyboardButton(
        text=f"📚 {t.title}",
        callback_data=f"admin:topic:{t.id}"
    )] for t in topics]

    await message.answer("📋 Testlar ro'yxati:", reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons))


@router.callback_query(F.data.startswith("admin:topic:"))
async def topic_detail(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        return

    topic_id = int(callback.data.split(":")[2])

    async with Session() as session:
        topic = await session.get(Topic, topic_id)
        naz_count = await session.scalar(
            select(func.count()).select_from(Question).where(
                Question.topic_id == topic_id, Question.type == "nazariy"
            )
        )
        aml_count = await session.scalar(
            select(func.count()).select_from(Question).where(
                Question.topic_id == topic_id, Question.type == "amaliy"
            )
        )

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🗑 O'chirish", callback_data=f"admin:delete:{topic_id}")],
        [InlineKeyboardButton(text="🔙 Orqaga",    callback_data="admin:back_tests")],
    ])

    await callback.message.answer(
        f"📚 {topic.title}\n"
        f"📖 Nazariy savollar: {naz_count} ta\n"
        f"🔧 Amaliy savollar: {aml_count} ta",
        reply_markup=keyboard,
    )


@router.callback_query(F.data == "admin:back_tests")
async def back_to_tests(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        return
    await _show_tests_list(callback.message)


@router.callback_query(F.data.startswith("admin:delete:"))
async def delete_topic(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        return

    topic_id = int(callback.data.split(":")[2])

    async with Session() as session:
        topic = await session.get(Topic, topic_id)
        topic.is_active = False
        await session.commit()

    await callback.message.answer(f"🗑 Test o'chirildi (arxivlandi).")


# ── Natijalar ─────────────────────────────────────────────────────────────────

async def _show_results(message: Message):
    async with Session() as session:
        results = (await session.scalars(
            select(UserResult).order_by(UserResult.passed_at.desc()).limit(50)
        )).all()

        if not results:
            await message.answer("📭 Hech qanday natija yo'q.")
            return

        user_ids = {r.user_id for r in results}
        topic_ids = {r.topic_id for r in results}
        users  = {u.id: u for u in (await session.scalars(select(User).where(User.id.in_(user_ids)))).all()}
        topics = {t.id: t for t in (await session.scalars(select(Topic).where(Topic.id.in_(topic_ids)))).all()}

    lines = ["📊 So'nggi natijalar:\n"]
    for r in results[:20]:
        u = users.get(r.user_id)
        t = topics.get(r.topic_id)
        icon = "✅" if r.score >= 80 else "❌"
        lines.append(f"{icon} {u.name} — {t.title} ({r.type}): {r.score:.0f}%")

    await message.answer("\n".join(lines))


# ── Word eksport ──────────────────────────────────────────────────────────────

async def _export_word(message: Message):
    async with Session() as session:
        results = (await session.scalars(select(UserResult).order_by(UserResult.score.desc()))).all()
        user_ids  = {r.user_id for r in results}
        topic_ids = {r.topic_id for r in results}
        users  = {u.id: u for u in (await session.scalars(select(User).where(User.id.in_(user_ids)))).all()}
        topics = {t.id: t for t in (await session.scalars(select(Topic).where(Topic.id.in_(topic_ids)))).all()}

    doc = Document()
    doc.add_heading("Natijalar hisoboti", 0)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Ism", "Dars", "Tur", "Ball", "Sana"

    for r in results:
        u = users.get(r.user_id)
        t = topics.get(r.topic_id)
        row = table.add_row().cells
        row[0].text = u.name if u else "—"
        row[1].text = t.title if t else "—"
        row[2].text = r.type
        row[3].text = f"{r.score:.0f}%"
        row[4].text = r.passed_at.strftime("%d.%m.%Y %H:%M")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    await message.answer_document(
        BufferedInputFile(buf.read(), filename="natijalar.docx"),
        caption="📥 Natijalar hisoboti",
    )
