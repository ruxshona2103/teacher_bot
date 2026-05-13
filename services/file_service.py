import re
import docx
import fitz
from pathlib import Path

# Kirill → Lotin harflar normalizatsiya
_CYRILLIC_MAP = {'А': 'A', 'В': 'B', 'С': 'C', 'Д': 'D',
                 'а': 'a', 'в': 'b', 'с': 'c', 'д': 'd'}


def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".docx":
        return _read_docx(file_path)
    elif ext == ".pdf":
        return _read_pdf(file_path)
    elif ext == ".txt":
        return _read_txt(file_path)
    else:
        raise ValueError(f"Qo'llab-quvvatlanmaydigan format: {ext}")


def _read_docx(path: str) -> str:
    doc = docx.Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                tables_text.append(row_text)
    return "\n".join(paragraphs + tables_text)


def _read_pdf(path: str) -> str:
    doc = fitz.open(path)
    pages = [page.get_text() for page in doc]
    doc.close()
    return "\n".join(pages)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def count_expected_questions(text: str) -> int:
    return len(re.findall(r'(?m)^\d+[.)]\s', text))


def parse_questions_with_regex(text: str) -> list[dict]:
    questions = []
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Har bir savol blokini ajratib olish
    blocks = re.split(r'\n(?=\d+[.)]\s)', text.strip())

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        lines = [l.strip() for l in block.split('\n') if l.strip()]
        if len(lines) < 5:
            continue

        # Birinchi qator: raqam + savol matni
        q_lines = [re.sub(r'^\d+[.)]\s*', '', lines[0])]

        # Variant boshlanishini toping
        option_start = None
        for i, line in enumerate(lines[1:], 1):
            norm = _normalize_letter(line[0]) if line else ''
            if re.match(r'^[A-Da-d][.)]\s', _normalize_letter(line)):
                option_start = i
                break
            q_lines.append(line)

        if option_start is None:
            continue

        q_text = ' '.join(q_lines).strip()

        # Variantlarni parse qilish
        option_map = {}
        for line in lines[option_start:]:
            norm_line = _normalize_letter(line)
            m = re.match(r'^([A-Da-d])[.)]\s*(.*)', norm_line)
            if m:
                letter = m.group(1).upper()
                if letter not in option_map:
                    option_map[letter] = line[2:].strip() if len(line) > 2 else ''

        if not all(k in option_map for k in ['A', 'B', 'C', 'D']):
            continue

        questions.append({
            'text': q_text,
            'option_a': option_map['A'],
            'option_b': option_map['B'],
            'option_c': option_map['C'],
            'option_d': option_map['D'],
        })

    return questions


def _normalize_letter(text: str) -> str:
    return ''.join(_CYRILLIC_MAP.get(c, c) for c in text)


def is_standard_format(text: str) -> bool:
    return count_expected_questions(text) >= 3
